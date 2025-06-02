# Bibliotecas
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Importa variáveis da interface
from UI import prod_codigo
from UI import nome_produto

# Cria documento Word
doc = Document()

# Define as configurações da página: 2x3 polegadas e sem margens
section = doc.sections[0]
section.page_width = Inches(2)
section.page_height = Inches(3)
section.top_margin = section.right_margin = section.left_margin = section.bottom_margin = 0

# Adiciona título e código do produto
title = doc.add_heading(f"Etiqueta: {nome_produto}", level=1)
text = doc.add_paragraph(str(prod_codigo))

# Insere QR Code da imagem gerada
qrcode_p = doc.add_paragraph("")
run = qrcode_p.add_run()
run.add_picture(f"qrcode_{prod_codigo}.png", width=Inches(1.0))

# Centraliza todos os elementos
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
text.alignment = WD_ALIGN_PARAGRAPH.CENTER
qrcode_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Salva o documento na pasta Docs/
filename = f"Docs/tag_{prod_codigo}.docx"
os.makedirs(os.path.dirname(filename), exist_ok=True)
doc.save(filename)
print(f"Documento salvo como: {os.path.abspath(filename)}")
